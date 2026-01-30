from __future__ import annotations
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .parse_cv import CVModel, topo_order

def _add_title(doc: Document, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)

def _add_bullet(doc: Document, text: str):
    # Use the built-in List Bullet style to avoid any odd characters
    doc.add_paragraph(text, style="List Bullet")

def render_docx(model: CVModel, output_path: str, title: str | None = None):
    doc = Document()
    _add_title(doc, title or f"Rebuild Guide — {model.cv_id} → SAP Datasphere")

    # Context
    _add_heading(doc, "Context", 1)
    doc.add_paragraph(
        "This guide was generated from a HANA Calculation View design-time artifact (.hdbcalculationview). "
        "Calculation Views are XML design-time files compiled by HDI at deploy time. "
        "In SAP Business Data Cloud (BDC), Data Products are created in SAP Datasphere; "
        "these steps rebuild the logic natively there."
    )

    # Parameters
    if model.parameters:
        _add_heading(doc, "Parameters to define in Datasphere", 1)
        for p in model.parameters:
            _add_bullet(doc, f"{p['id']} ({p['sqlType']}), default '{p['defaultValue']}', mandatory: {p['isMandatory']}")

    # Data sources
    if model.data_sources:
        _add_heading(doc, "Source objects (prepare as Remote/Replicated Tables)", 1)
        for ds_id, uri in model.data_sources.items():
            _add_bullet(doc, f"{ds_id} -> {uri}")

    # Steps
    _add_heading(doc, "Step-by-step modeling in Datasphere", 1)
    order = topo_order(model)
    step = 1
    for node_id in order:
        node = model.nodes[node_id]
        _add_heading(doc, f"{step}. Create view for node '{node.node_id}' ({node.node_type})", 2)

        if node.node_type == "ProjectionView":
            _add_bullet(doc, "Create a Graphical View")
            if node.inputs:
                _add_bullet(doc, f"Use source: {node.inputs[0]}")
            if node.mappings:
                # list selected columns from mappings
                cols = sorted({m.target for m in node.mappings if m.target})
                if cols:
                    _add_bullet(doc, "Select columns: " + ", ".join(cols))
            for flt in node.filters:
                _add_bullet(doc, f"Add filter: {flt}")

        elif node.node_type == "JoinView":
            _add_bullet(doc, "Create a Graphical View with a Join node")
            if node.join_type:
                _add_bullet(doc, f"Join type: {node.join_type}")
            if node.inputs:
                _add_bullet(doc, "Inputs: " + ", ".join(node.inputs))
            if node.join_condition:
                _add_bullet(doc, f"Join condition: {node.join_condition}")

        elif node.node_type == "AggregationView":
            _add_bullet(doc, "Create a Graphical View with an Aggregation node")
            if node.attributes:
                _add_bullet(doc, "Group by attributes: " + ", ".join(node.attributes))
            if node.measures:
                _add_bullet(doc, "Define measures: " + ", ".join(node.measures))
            if node.calculated_measures:
                _add_bullet(doc, "Add calculated measures:")
                for mid, fml in node.calculated_measures.items():
                    _add_bullet(doc, f"  - {mid} = {fml}")
            if node.mappings:
                _add_bullet(doc, "Map inputs to output columns:")
                for mp in node.mappings:
                    _add_bullet(doc, f"  - {mp.source} -> {mp.target}")

        elif node.node_type == "UnionView":
            _add_bullet(doc, "Create a Graphical View with a Union node")
            if node.inputs:
                _add_bullet(doc, "Union inputs: " + ", ".join(node.inputs))
            _add_bullet(doc, "Align columns across inputs by name and type")

        else:
            _add_bullet(doc, "Node type not fully handled in MVP — add manual steps here.")

        step += 1

    # Semantics
    _add_heading(doc, "Business Semantics (Datasphere)", 1)
    if model.logical_attributes:
        _add_bullet(doc, "Mark as Attributes: " + ", ".join(model.logical_attributes))
    if model.logical_measures:
        _add_bullet(doc, "Mark as Measures: " + ", ".join(model.logical_measures))

    # Validation
    _add_heading(doc, "Validation", 1)
    _add_bullet(doc, "Compare row counts with the original Calculation View for a known time slice.")
    _add_bullet(doc, "Spot-check joins and calculated measures with a sample of document numbers.")

    # Publish
    _add_heading(doc, "Publish as a BDC Data Product", 1)
    _add_bullet(doc, "Package the final Datasphere view into a Data Product and add owners/tags/description.")
    _add_bullet(doc, "Expose via BDC Connect / Delta Sharing for Azure consumption, if required.")

    # Known gaps
    _add_heading(doc, "Known gaps / Manual actions", 1)
    _add_bullet(doc, "If the original view used SQLScript/table functions, create a Datasphere SQL View and port logic.")
    _add_bullet(doc, "Re-implement analytic privileges/row-level filters in Datasphere.")
    _add_bullet(doc, "Validate currency conversions and text-join semantics.")

    doc.save(output_path)