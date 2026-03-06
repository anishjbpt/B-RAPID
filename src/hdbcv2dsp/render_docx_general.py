# render_docx_general.py (enhanced for ABAP CDS)
from __future__ import annotations
from typing import Dict, List, Optional, Tuple
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Project imports
from .parse_cv import CVModel, topo_order as cv_topo_order
from .parse_sql_view import SQLViewModel
from .parse_procedure import ProcedureModel
from .parse_abap_cds import ABAPCDSModel  # NEW
from .artifacts import ArtifactNode, topo_order_nodes
from .summarize import summarize_cv, summarize_sql_view, summarize_procedure, summarize_abap_cds  # NEW

#############################
# Formatting helpers
#############################
def _title(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)

def _bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")

def _subbullet(doc: Document, text: str):
    # Indented bullet via a leading hyphen; Word will keep same bullet style
    doc.add_paragraph(f"- {text}", style="List Bullet")


def _step_by_step_cv_in_datasphere(doc: Document):
    _heading(doc, "Step-by-step set-up in Datasphere", 2)
    _bullet(doc, "Prepare **sources**: import as **Remote Tables** (federate) or use **Replication Flows**/**Remote Table replication** if you need persisted/local data.")
    _bullet(doc, "Rebuild the logic:")
    _subbullet(doc, "Option A (SQL): create a **SQL View** and translate node logic (Projection/Join/Union/Aggregation) into SQL; define **input parameters** (if any) in the editor.")
    _subbullet(doc, "Option B (Graphical): create a **Graphical View**, add sources, and recreate joins/aggregations and calculated columns.")
    _bullet(doc, "Validate:")
    _subbullet(doc, "Compare record counts for a known slice; reconcile **joins**, **filters**, and **measures**.")
    _subbullet(doc, "Check **semantic settings** (data category, view type) and descriptions.")
    _bullet(doc, "Deploy the view and smoke‑test with business keys.")

def _notes_cv(doc: Document):
    _heading(doc, "Notes & Considerations", 3)
    _bullet(doc, "If the original CV used a **star join**, prefer a single SQL View with explicit joins or a Graphical View with a star‑join layout.")
    _bullet(doc, "Recreate **calculated measures** and **filters** exactly; confirm numeric precision/scale and date/timestamp semantics.")
    _bullet(doc, "If parameters/variables exist, define them in the Datasphere view and reflect them in SQL filters or input mappings.")

def _step_by_step_sql_in_datasphere(doc: Document):
    _heading(doc, "Step-by-step set-up in Datasphere", 2)
    _bullet(doc, "Ensure **base tables/views** exist in the target space (Remote or Local/Replicated).")
    _bullet(doc, "Create an **SQL View** and paste the final SQL.")
    _bullet(doc, "Adjust **column data types** where required and set labels/semantics.")
    _bullet(doc, "Deploy and validate using representative business keys; reconcile aggregations and filters.")

def _notes_sql(doc: Document):
    _heading(doc, "Notes & Considerations", 3)
    _bullet(doc, "If your SQL uses database‑specific functions, check dialect compatibility in Datasphere and replace with supported expressions.")
    _bullet(doc, "For very large joins/aggregations, consider **replicating** hot tables to local storage for performance.")
    _bullet(doc, "Document any hard‑coded date ranges / predicates and externalize them via input parameters if needed.")

def _step_by_step_proc_in_datasphere(doc: Document):
    _heading(doc, "Step-by-step set-up in Datasphere", 2)
    _bullet(doc, "Identify the **stages** inside the procedure (reads → transforms → writes).")
    _bullet(doc, "Rebuild as **SQL Views** and/or **Transformation Flows**:")
    _subbullet(doc, "Set up **Local Tables** for intermediate staging if needed.")
    _subbullet(doc, "Translate procedural steps into SQL transformations or flow operators (projection, join, filter, aggregation).")
    _bullet(doc, "If the procedure writes to permanent targets, create those **Local Tables** first.")
    _bullet(doc, "Deploy and run end‑to‑end; validate row counts and key‑by‑key samples.")

def _notes_proc(doc: Document):
    _heading(doc, "Notes & Considerations", 3)
    _bullet(doc, "Break monolithic logic into smaller, testable transformations; avoid temp‑table name clashes by using deterministic table/view names.")
    _bullet(doc, "Parameterize constants (dates, thresholds) via Task Chains/CLI, or convert them to view parameters.")
    _bullet(doc, "If the original procedure called other procedures, inline or modularize them as separate views/flows and orchestrate with **Task Chains**.")




def _fmt_list(items: List[str], limit: int | None = None) -> str:
    if not items:
        return ""
    if limit and len(items) > limit:
        return ", ".join(items[:limit]) + f" … (+{len(items)-limit})"
    return ", ".join(items)

#############################
# SQL text analyzers (heuristic)
#############################
# (unchanged helper functions omitted for brevity...)

# ... keep all existing helpers here unchanged ...

#############################
# Step-by-step builders
#############################

def _step_by_step_sql_view(doc: Document, v: SQLViewModel):
    # (existing body unchanged)
    # ...

    # NOTE: body truncated here in this snippet for brevity, keep your full original function
    pass  # <-- remove this pass; keep your original function body

def _step_by_step_calc_view(doc: Document, cv: CVModel):
    # (existing body unchanged)
    pass  # <-- remove this pass; keep your original function body

def _step_by_step_procedure(doc: Document, p: ProcedureModel):
    # (existing body unchanged)
    pass  # <-- remove this pass; keep your original function body

# NEW: ABAP CDS replication guidance
def _step_by_step_abap_cds(doc: Document, cds: ABAPCDSModel):
    _heading(doc, "Step-by-step set-up in Datasphere (Replication Flow)", 2)
    _bullet(doc, "Create a **Replication Flow**:")
    _subbullet(doc, "Source connection: ABAP/SAP S/4HANA; Container: *CDS Views Enabled for Data Extraction*.")
    _subbullet(doc, f"Add source object: `{cds.name}`.")
    _subbullet(doc, "Target: Local table (new or map to existing). Choose **Load Type**: *Initial and Delta* when CDC is available.")
    _subbullet(doc, "Optionally tune **Source/Target Thread Limits** and schedule delta frequency.")
    _bullet(doc, "Deploy and **Run** the replication flow, then validate row counts against the CDS view output.")
    # Hints
    _heading(doc, "Notes & Considerations", 3)
    _bullet(doc, "Replication Flows don’t support input parameters — parameterized CDS entities can’t be replicated as-is.")
    _bullet(doc, "Ensure `@Analytics.dataExtraction.enabled: true`; add CDC mapping for delta where applicable.")
    if cds.parameters:
        _bullet(doc, "This CDS has parameters → consider creating a specialized CDS without parameters, or consume via remote table/SQL view.")

#############################
# Main renderer
#############################
def render_docx_general(
    output_path: str,
    title: Optional[str],
    cv_model: Optional[CVModel] = None,
    sql_views: Optional[List[SQLViewModel]] = None,
    procedures: Optional[List[ProcedureModel]] = None,
    graph: Optional[Dict[str, ArtifactNode]] = None,
    abap_cds_list: Optional[List[ABAPCDSModel]] = None,  # NEW
):
    """
    Renders a mixed-artifact DOCX guide with a consistent structure across:
    - HANA Calculation Views
    - SQL Views
    - Stored Procedures
    - ABAP CDS (NEW)
    """
    sql_views = sql_views or []
    procedures = procedures or []
    abap_cds_list = abap_cds_list or []

    doc = Document()
    _title(doc, title or "Rebuild Guide — SAP HANA Artifacts → SAP Datasphere")

    # Context
    _heading(doc, "Context", 1)
    doc.add_paragraph(
        "This guide consolidates design-time artifacts (HANA Calculation Views, SQL Views, Stored Procedures, and ABAP CDS) "
        "and outlines practical steps to rebuild the logic in SAP Datasphere / Business Data Cloud."
    )

    # -----------------------------
    # ABAP CDS section (NEW)
    # -----------------------------
    if abap_cds_list:
        _heading(doc, "ABAP CDS", 1)
        for cds in abap_cds_list:
            _heading(doc, f"CDS: {cds.name}", 2)
            _heading(doc, "Understanding (plain-English)", 3)
            for s in summarize_abap_cds(cds):
                _bullet(doc, s)
            _step_by_step_abap_cds(doc, cds)

    # -----------------------------
    # Calculation View section
    # -----------------------------
    if cv_model:
        _heading(doc, f"Calculation View: {cv_model.cv_id}", 1)
        _heading(doc, "Understanding (plain-English)", 2)
        for s in summarize_cv(cv_model):
            _bullet(doc, s)
        if getattr(cv_model, 'description', None):
            _bullet(doc, f"Description: {cv_model.description}")
        _bullet(doc, f"Output View Type: {getattr(cv_model, 'output_view_type', '')} • "
                     f"Data Category: {getattr(cv_model, 'data_category', '')}")
        if getattr(cv_model, 'parameters', None):
            _heading(doc, "Parameters (define in Datasphere)", 2)
            for p in cv_model.parameters:
                _bullet(doc, f"{p['id']} ({p['sqlType']}), default '{p['defaultValue']}', mandatory: {p['isMandatory']}")
        if getattr(cv_model, 'data_sources', None):
            _heading(doc, "Source objects (prepare as Remote/Replicated Tables)", 2)
            for ds_id, uri in cv_model.data_sources.items():
                _bullet(doc, f"{ds_id} → {uri}")
        _step_by_step_calc_view(doc, cv_model)


        # === NEW: tailored steps + notes for Calculation View ===
        _step_by_step_cv_in_datasphere(doc)
        _notes_cv(doc)



    # -----------------------------
    # SQL Views section
    # -----------------------------
    if sql_views:
        _heading(doc, "SQL Views", 1)
        for v in sql_views:
            _heading(doc, f"SQL View: {v.name}", 2)
            _heading(doc, "Understanding (plain-English)", 3)
            for s in summarize_sql_view(v):
                _bullet(doc, s)
            if getattr(v, 'inputs', None):
                _bullet(doc, "Upstream sources: " + _fmt_list(v.inputs))
            if getattr(v, 'columns', None):
                preview = _fmt_list(v.columns, limit=10)
                _bullet(doc, "Output columns (preview): " + preview)
            _step_by_step_sql_view(doc, v)


            # === NEW: tailored steps + notes for SQL View ===
            _step_by_step_sql_in_datasphere(doc)
            _notes_sql(doc)
          

    # -----------------------------
    # Procedures section
    # -----------------------------
    if procedures:
        _heading(doc, "Stored Procedures", 1)
        for p in procedures:
            _heading(doc, f"Procedure: {p.name}", 2)
            _heading(doc, "Understanding (plain-English)", 3)
            for s in summarize_procedure(p):
                _bullet(doc, s)
            if getattr(p, 'parameters', None):
                _bullet(doc, "Parameters: " + ", ".join([f"{x['mode']} {x['name']} {x['type']}" for x in p.parameters]))
            if getattr(p, 'reads_from', None):
                _bullet(doc, "Reads from: " + _fmt_list(p.reads_from))
            if getattr(p, 'writes_to', None):
                _bullet(doc, "Writes to (permanent): " + _fmt_list(p.writes_to))
            if getattr(p, 'temp_tables', None):
                _bullet(doc, "Temp tables used: " + _fmt_list(p.temp_tables))
            _step_by_step_procedure(doc, p)


            # === NEW: tailored steps + notes for Procedure ===
            _step_by_step_proc_in_datasphere(doc)
            _notes_proc(doc)



    # -----------------------------
    # Combined dependency order (optional)
    # -----------------------------
    if graph:
        _heading(doc, "Combined Dependency Order (best effort)", 1)
        order = topo_order_nodes(graph)
        for i, node_id in enumerate(order, 1):
            node = graph[node_id]
            _bullet(doc, f"{i}. {node_id} ({node.kind})")

    # Validation & Publish
    _heading(doc, "Validation", 1)
    _bullet(doc, "Compare row counts against the source artifacts for a known time slice.")
    _bullet(doc, "Spot-check joins and calculations using representative business keys.")
    _heading(doc, "Publish as a BDC Data Product", 1)
    _bullet(doc, "Package the final Datasphere view into a Data Product with owners/tags/description.")

    # Save
    doc.save(output_path)